#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
dokument_pipeline.py — Kira Dokumente: Import-Pipeline (session-eeee)

Verantwortlich für:
  1. Datei-Import (Ordner, Upload, Mail-Anhang)
  2. Text-Extraktion (PDF, DOCX, Bild-OCR)
  3. LLM-Klassifizierung
  4. Routing-Entscheidung (erfordert_handlung, routing_ziel, zielmodul)
  5. Vorgangszuordnung
  6. Runtime-Log-Events
"""
import json
import logging
import mimetypes
from pathlib import Path
from datetime import datetime

SCRIPTS_DIR   = Path(__file__).parent
KNOWLEDGE_DIR = SCRIPTS_DIR.parent / "knowledge"
CONFIG_FILE   = SCRIPTS_DIR / "config.json"

log = logging.getLogger("dokument_pipeline")

# ── Lazy Imports für optionale Dependencies ───────────────────────────────────

def _has_pdfplumber():
    try:
        import pdfplumber
        return True
    except ImportError:
        return False

def _has_pytesseract():
    try:
        import pytesseract
        return True
    except ImportError:
        return False

def _has_docx():
    try:
        import docx
        return True
    except ImportError:
        return False

# ── Konfiguration ─────────────────────────────────────────────────────────────

def _load_config() -> dict:
    try:
        return json.loads(CONFIG_FILE.read_text(encoding="utf-8"))
    except Exception:
        return {}

def _get_dok_config() -> dict:
    return _load_config().get("dokumente", {})

# ── Text-Extraktion ──────────────────────────────────────────────────────────

def extract_text_pdf(file_path: Path) -> str:
    """Extrahiert Text aus PDF via pdfplumber."""
    if not _has_pdfplumber():
        log.warning("pdfplumber nicht installiert — PDF-Text kann nicht extrahiert werden")
        return ""
    import pdfplumber
    text_parts = []
    try:
        with pdfplumber.open(str(file_path)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
    except Exception as e:
        log.error(f"PDF-Extraktion fehlgeschlagen: {e}")
    return "\n".join(text_parts)


def extract_text_docx(file_path: Path) -> str:
    """Extrahiert Text aus DOCX via python-docx."""
    if not _has_docx():
        log.warning("python-docx nicht installiert — DOCX-Text kann nicht extrahiert werden")
        return ""
    import docx
    try:
        doc = docx.Document(str(file_path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        log.error(f"DOCX-Extraktion fehlgeschlagen: {e}")
        return ""


def extract_text_image(file_path: Path) -> str:
    """OCR für Bilder via pytesseract."""
    if not _has_pytesseract():
        log.warning("pytesseract nicht installiert — Bild-OCR nicht verfügbar")
        return ""
    import pytesseract
    try:
        from PIL import Image
        img = Image.open(str(file_path))
        text = pytesseract.image_to_string(img, lang="deu+eng")
        return text.strip()
    except Exception as e:
        log.error(f"OCR fehlgeschlagen: {e}")
        return ""


def extract_text(file_path: Path) -> str:
    """Erkennt Dateityp und extrahiert Text."""
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        text = extract_text_pdf(file_path)
        if not text.strip() and _has_pytesseract():
            # PDF ohne Text → OCR-Versuch
            log.info(f"PDF ohne Text, versuche OCR: {file_path.name}")
            text = _ocr_pdf_pages(file_path)
        return text
    elif suffix in (".docx", ".doc"):
        return extract_text_docx(file_path)
    elif suffix in (".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".webp"):
        return extract_text_image(file_path)
    elif suffix in (".txt", ".md", ".csv"):
        try:
            return file_path.read_text(encoding="utf-8", errors="replace")
        except Exception:
            return ""
    return ""


def _ocr_pdf_pages(file_path: Path) -> str:
    """OCR für gescannte PDF-Seiten (Bilder extrahieren und OCR)."""
    if not _has_pdfplumber() or not _has_pytesseract():
        return ""
    import pdfplumber
    import pytesseract
    from PIL import Image
    import io
    text_parts = []
    try:
        with pdfplumber.open(str(file_path)) as pdf:
            for page in pdf.pages[:20]:  # Max 20 Seiten
                img = page.to_image(resolution=200)
                pil_img = img.original
                t = pytesseract.image_to_string(pil_img, lang="deu+eng")
                if t.strip():
                    text_parts.append(t.strip())
    except Exception as e:
        log.error(f"PDF-OCR fehlgeschlagen: {e}")
    return "\n".join(text_parts)


# ── Dateityp-Erkennung ───────────────────────────────────────────────────────

SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".doc", ".xlsx", ".xls",
    ".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".webp",
    ".txt", ".md", ".csv", ".html", ".htm",
}

def is_supported(file_path: Path) -> bool:
    return file_path.suffix.lower() in SUPPORTED_EXTENSIONS


def detect_dateityp(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    mapping = {
        ".pdf": "pdf", ".docx": "docx", ".doc": "doc",
        ".xlsx": "xlsx", ".xls": "xls",
        ".jpg": "bild", ".jpeg": "bild", ".png": "bild",
        ".tiff": "bild", ".tif": "bild", ".bmp": "bild", ".webp": "bild",
        ".txt": "text", ".md": "text", ".csv": "tabelle",
        ".html": "html", ".htm": "html",
    }
    return mapping.get(suffix, "sonstig")


# ── LLM-Klassifizierung ─────────────────────────────────────────────────────

ROUTING_ZIELE = ["task", "buchhaltung", "feed", "kira_vorschlag", "archivieren", "manuelle_pruefung"]

def classify_dokument(text: str, dateiname: str, dateityp: str) -> dict:
    """
    Klassifiziert ein Dokument via LLM.
    Gibt zurück: {kategorie, dokumentrolle, erfordert_handlung, routing_ziel, konfidenz, begruendung}
    """
    if not text.strip():
        return {
            "kategorie": "unbekannt",
            "dokumentrolle": "unbekannt",
            "erfordert_handlung": False,
            "routing_ziel": "manuelle_pruefung",
            "konfidenz": 0.0,
            "begruendung": "Kein Text extrahiert",
        }

    try:
        from kira_llm import get_provider_for_task, _call_provider
    except ImportError:
        log.error("kira_llm nicht verfügbar — Fallback auf regelbasierte Klassifizierung")
        return _classify_regelbasiert(text, dateiname, dateityp)

    prompt = f"""Analysiere dieses Dokument und klassifiziere es.

Dateiname: {dateiname}
Dateityp: {dateityp}
Text (erste 3000 Zeichen):
{text[:3000]}

Antworte AUSSCHLIESSLICH als JSON mit diesen Feldern:
{{
  "kategorie": "rechnung|angebot|mahnung|zahlungserinnerung|vertrag|brief|behoerdenpost|versicherung|mitarbeiterdokument|notiz|auswertung|sonstiges",
  "dokumentrolle": "eingangsrechnung|ausgangsrechnung|eingangsangebot|ausgangsangebot|mahnung_eingang|mahnung_ausgang|vertrag|korrespondenz|behoerdlich|intern|referenz|beleg|sonstiges",
  "erfordert_handlung": true/false,
  "routing_ziel": "task|buchhaltung|feed|kira_vorschlag|archivieren|manuelle_pruefung",
  "konfidenz": 0.0-1.0,
  "begruendung": "kurze Begründung"
}}"""

    try:
        provider = get_provider_for_task("classify")
        if provider:
            result = _call_provider(
                provider["typ"], provider["url"], provider["api_key"],
                provider["model"], prompt, max_tokens=500,
            )
            import re
            json_match = re.search(r'\{[^{}]*\}', result, re.DOTALL)
            if json_match:
                parsed = json.loads(json_match.group())
                # Validierung
                if parsed.get("routing_ziel") not in ROUTING_ZIELE:
                    parsed["routing_ziel"] = "manuelle_pruefung"
                return parsed
    except Exception as e:
        log.error(f"LLM-Klassifizierung fehlgeschlagen: {e}")

    return _classify_regelbasiert(text, dateiname, dateityp)


def _classify_regelbasiert(text: str, dateiname: str, dateityp: str) -> dict:
    """Einfache regelbasierte Klassifizierung als Fallback."""
    text_lower = text.lower()
    name_lower = dateiname.lower()

    if any(w in text_lower for w in ["rechnung", "invoice", "rechnungsnummer", "re-"]):
        return {
            "kategorie": "rechnung",
            "dokumentrolle": "eingangsrechnung",
            "erfordert_handlung": True,
            "routing_ziel": "buchhaltung",
            "konfidenz": 0.6,
            "begruendung": "Enthält Rechnungs-Keywords",
        }
    if any(w in text_lower for w in ["angebot", "angebotsnummer", "offerte"]):
        return {
            "kategorie": "angebot",
            "dokumentrolle": "eingangsangebot",
            "erfordert_handlung": True,
            "routing_ziel": "task",
            "konfidenz": 0.6,
            "begruendung": "Enthält Angebots-Keywords",
        }
    if any(w in text_lower for w in ["mahnung", "zahlungserinnerung", "letzte aufforderung"]):
        return {
            "kategorie": "mahnung",
            "dokumentrolle": "mahnung_eingang",
            "erfordert_handlung": True,
            "routing_ziel": "task",
            "konfidenz": 0.7,
            "begruendung": "Enthält Mahnungs-Keywords",
        }
    if any(w in text_lower for w in ["vertrag", "vereinbarung", "unterzeichn"]):
        return {
            "kategorie": "vertrag",
            "dokumentrolle": "vertrag",
            "erfordert_handlung": True,
            "routing_ziel": "kira_vorschlag",
            "konfidenz": 0.5,
            "begruendung": "Enthält Vertrags-Keywords",
        }

    return {
        "kategorie": "sonstiges",
        "dokumentrolle": "sonstiges",
        "erfordert_handlung": False,
        "routing_ziel": "archivieren",
        "konfidenz": 0.3,
        "begruendung": "Keine spezifischen Keywords erkannt",
    }


# ── Routing ───────────────────────────────────────────────────────────────────

ROUTING_ZU_ZIELMODUL = {
    "task":               "kommunikation",
    "buchhaltung":        "geschaeft",
    "feed":               "dashboard",
    "kira_vorschlag":     "kira",
    "archivieren":        "dokumente",
    "manuelle_pruefung":  "dokumente",
}

def route_dokument(klassifizierung: dict) -> dict:
    """Bestimmt das Zielmodul basierend auf der Routing-Entscheidung."""
    routing_ziel = klassifizierung.get("routing_ziel", "manuelle_pruefung")
    zielmodul = ROUTING_ZU_ZIELMODUL.get(routing_ziel, "dokumente")
    return {
        "routing_ziel": routing_ziel,
        "zielmodul": zielmodul,
        "erfordert_handlung": klassifizierung.get("erfordert_handlung", False),
    }


# ── Vollständige Pipeline ────────────────────────────────────────────────────

def process_file(
    file_path: Path,
    quelle: str = "upload",
    quell_id: str = "",
    kunde_name: str = "",
    vorgang_id: int = None,
    move_file: bool = False,
) -> dict:
    """
    Vollständige Dokument-Pipeline:
    1. Dateityp erkennen
    2. Hash berechnen + Dedup prüfen
    3. Text extrahieren (PDF/DOCX/OCR)
    4. Klassifizieren (LLM oder regelbasiert)
    5. Routing-Entscheidung
    6. In externe Ablage speichern
    7. DB-Eintrag anlegen
    8. Runtime-Log schreiben

    Gibt das fertige Dokument-Dict zurück.
    """
    from dokument_storage import (
        compute_hash, check_duplicate, build_ablage_pfad,
        store_file, create_dokument,
    )

    if not file_path.exists():
        return {"error": f"Datei nicht gefunden: {file_path}"}

    if not is_supported(file_path):
        return {"error": f"Dateityp nicht unterstützt: {file_path.suffix}"}

    # 1. Grunddaten
    dateiname = file_path.name
    dateityp = detect_dateityp(file_path)
    dateigroesse = file_path.stat().st_size

    # 2. Hash + Dedup
    hash_val = compute_hash(file_path)
    existing = check_duplicate(hash_val)
    if existing:
        log.info(f"Dublette erkannt: {dateiname} → bestehend ID {existing['id']}")
        _log_event("doc_dedup", f"Dublette: {dateiname}", {
            "original_id": existing["id"], "hash": hash_val,
        })
        return {"duplicate": True, "existing_id": existing["id"], "dateiname": dateiname}

    # 3. Text extrahieren
    text = extract_text(file_path)
    log.info(f"Text extrahiert: {len(text)} Zeichen aus {dateiname}")

    # 4. Klassifizierung
    klassifizierung = classify_dokument(text, dateiname, dateityp)
    log.info(f"Klassifiziert: {dateiname} → {klassifizierung.get('kategorie')} "
             f"(Konfidenz: {klassifizierung.get('konfidenz', 0):.0%})")

    # 5. Routing
    routing = route_dokument(klassifizierung)

    # 6. Externe Ablage
    ziel_pfad = build_ablage_pfad(
        dateiname, quelle, kunde_name, vorgang_id,
        klassifizierung.get("dokumentrolle", ""),
    )
    store_file(file_path, ziel_pfad, move=move_file)

    # 7. DB-Eintrag
    dok_id = create_dokument(
        titel=file_path.stem,
        dateiname=dateiname,
        dateityp=dateityp,
        dateigroesse=dateigroesse,
        hash_sha256=hash_val,
        externer_pfad=str(ziel_pfad),
        quelle=quelle,
        quell_id=quell_id,
        dokumentrolle=klassifizierung.get("dokumentrolle", ""),
        kategorie=klassifizierung.get("kategorie", ""),
        ocr_text=text[:50000],  # Max 50k Zeichen OCR-Text in DB
        klassifizierung=klassifizierung,
        erfordert_handlung=klassifizierung.get("erfordert_handlung", False),
        routing_ziel=routing["routing_ziel"],
        zielmodul=routing["zielmodul"],
        vorgang_id=vorgang_id,
        konfidenz=klassifizierung.get("konfidenz", 0),
    )

    # 8. Runtime-Log
    _log_event("doc_import", f"Dokument importiert: {dateiname}", {
        "dok_id": dok_id, "quelle": quelle, "dateityp": dateityp,
        "groesse": dateigroesse, "text_laenge": len(text),
    })
    _log_event("doc_classify", f"Klassifiziert: {klassifizierung.get('kategorie')}", {
        "dok_id": dok_id, **klassifizierung,
    })
    _log_event("doc_route", f"Routing: {routing['routing_ziel']} → {routing['zielmodul']}", {
        "dok_id": dok_id, **routing,
    })

    result = {
        "id": dok_id,
        "dateiname": dateiname,
        "dateityp": dateityp,
        "dateigroesse": dateigroesse,
        "hash": hash_val,
        "externer_pfad": str(ziel_pfad),
        "klassifizierung": klassifizierung,
        "routing": routing,
        "text_laenge": len(text),
    }
    log.info(f"Pipeline abgeschlossen: {dateiname} → ID {dok_id}")
    return result


# ── Watched Folder ────────────────────────────────────────────────────────────

_watcher_thread = None

def start_folder_watcher():
    """Startet den Ordner-Watcher als Hintergrund-Thread."""
    global _watcher_thread
    if _watcher_thread and _watcher_thread.is_alive():
        return

    dok_cfg = _get_dok_config()
    if not dok_cfg.get("aktiv", False):
        log.info("Dokumente-Modul nicht aktiv — Watcher nicht gestartet")
        return

    watch_pfad = dok_cfg.get("ueberwachter_ordner", "")
    if not watch_pfad:
        log.info("Kein überwachter Ordner konfiguriert")
        return

    watch_dir = Path(watch_pfad)
    if not watch_dir.exists():
        log.warning(f"Überwachter Ordner existiert nicht: {watch_pfad}")
        return

    try:
        from watchdog.observers import Observer
        from watchdog.events import FileSystemEventHandler
    except ImportError:
        log.warning("watchdog nicht installiert — Ordner-Überwachung nicht verfügbar")
        return

    import threading
    import time

    class DokumentHandler(FileSystemEventHandler):
        def __init__(self):
            self._debounce = {}
            self._debounce_secs = 1.0

        def on_created(self, event):
            if event.is_directory:
                return
            fp = Path(event.src_path)
            if not is_supported(fp):
                return
            # Debounce: Windows feuert oft doppelte Events
            now = time.time()
            if fp in self._debounce and now - self._debounce[fp] < self._debounce_secs:
                return
            self._debounce[fp] = now
            # Kurz warten bis Datei vollständig geschrieben
            time.sleep(0.5)
            try:
                log.info(f"Neues Dokument erkannt: {fp.name}")
                process_file(fp, quelle="watched_folder", move_file=False)
            except Exception as e:
                log.error(f"Fehler bei Verarbeitung von {fp.name}: {e}")

    def _run_watcher():
        observer = Observer()
        rekursiv = dok_cfg.get("unterordner_ueberwachen", True)
        observer.schedule(DokumentHandler(), str(watch_dir), recursive=rekursiv)
        observer.start()
        log.info(f"Ordner-Watcher gestartet: {watch_dir} (rekursiv={rekursiv})")
        try:
            while True:
                time.sleep(1)
        except Exception:
            observer.stop()
        observer.join()

    _watcher_thread = threading.Thread(target=_run_watcher, daemon=True, name="dokument-watcher")
    _watcher_thread.start()


# ── Runtime-Log Helper ───────────────────────────────────────────────────────

def _log_event(action: str, summary: str, payload: dict = None):
    try:
        from runtime_log import elog
        elog("dokument", action, summary,
             modul="dokumente",
             vollkontext=payload)
    except Exception:
        log.debug(f"Runtime-Log nicht verfügbar: {action} — {summary}")


# ── Dokument-Export (PDF / DOCX) ─────────────────────────────────────────────

def export_pdf(html_content: str, briefkopf: dict = None) -> bytes:
    """Erzeugt PDF aus HTML via weasyprint."""
    try:
        import weasyprint
    except ImportError:
        raise RuntimeError("weasyprint nicht installiert — PDF-Export nicht verfügbar")

    # Briefkopf einbauen
    full_html = _build_full_html(html_content, briefkopf)
    pdf_bytes = weasyprint.HTML(string=full_html).write_pdf()
    return pdf_bytes


def export_docx_from_template(
    vorlage_pfad: Path,
    platzhalter: dict,
) -> bytes:
    """Erzeugt DOCX aus Vorlage via docxtpl."""
    try:
        from docxtpl import DocxTemplate
    except ImportError:
        raise RuntimeError("docxtpl nicht installiert — DOCX-Export nicht verfügbar")

    doc = DocxTemplate(str(vorlage_pfad))
    doc.render(platzhalter)
    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_docx_from_html(html_content: str, briefkopf: dict = None) -> bytes:
    """Erzeugt einfaches DOCX aus HTML-Inhalt."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise RuntimeError("python-docx nicht installiert — DOCX-Export nicht verfügbar")

    import re
    # Einfache HTML→Text-Konvertierung
    text = re.sub(r'<br\s*/?>', '\n', html_content)
    text = re.sub(r'<[^>]+>', '', text)
    text = text.strip()

    doc = Document()
    if briefkopf:
        if briefkopf.get("firmenname"):
            doc.add_heading(briefkopf["firmenname"], level=1)
        if briefkopf.get("adresse"):
            doc.add_paragraph(briefkopf["adresse"])

    for line in text.split("\n"):
        doc.add_paragraph(line)

    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_full_html(content: str, briefkopf: dict = None) -> str:
    """Baut vollständiges HTML mit Briefkopf für PDF-Export."""
    header = ""
    footer = ""
    if briefkopf:
        header = briefkopf.get("html_header", "")
        footer = briefkopf.get("html_footer", "")

    return f"""<!DOCTYPE html>
<html lang="de">
<head>
<meta charset="utf-8">
<style>
@page {{
    size: A4;
    margin: 2cm 2cm 3cm 2cm;
    @top-center {{ content: ""; }}
    @bottom-center {{ content: counter(page) " / " counter(pages); font-size: 9pt; color: #888; }}
}}
body {{
    font-family: 'Segoe UI', Arial, sans-serif;
    font-size: 11pt;
    line-height: 1.5;
    color: #222;
}}
h1 {{ font-size: 16pt; margin-bottom: 0.3em; }}
h2 {{ font-size: 13pt; margin-bottom: 0.2em; }}
table {{ border-collapse: collapse; width: 100%; }}
td, th {{ border: 1px solid #ccc; padding: 6px 10px; }}
.briefkopf {{ margin-bottom: 2em; border-bottom: 2px solid #333; padding-bottom: 1em; }}
.fussnote {{ margin-top: 2em; border-top: 1px solid #ccc; padding-top: 0.5em; font-size: 9pt; color: #666; }}
</style>
</head>
<body>
{f'<div class="briefkopf">{header}</div>' if header else ''}
{content}
{f'<div class="fussnote">{footer}</div>' if footer else ''}
</body>
</html>"""
